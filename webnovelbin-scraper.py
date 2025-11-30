import requests
from bs4 import BeautifulSoup
import time
from docx import Document
from urllib.parse import urljoin

# --- CONFIG ---
HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"
}
RETRY_LIMIT = 3
DELAY = 2


def fetch_page(url):
    """Safe request with retry"""
    for attempt in range(RETRY_LIMIT):
        try:
            res = requests.get(url, headers=HEADERS, timeout=10)
            if res.status_code == 200:
                return res.text
        except:
            print(f"Retry {attempt+1}/{RETRY_LIMIT} ...")
            time.sleep(2)
    return None


def clean_text(text):
    """Bersihkan newline berlebih & whitespace"""
    text = text.replace('\r', '').strip()
    lines = [line.strip() for line in text.split("\n")]
    lines = [ln for ln in lines if ln]   # remove empty line
    return lines


def scrape_chapter(url, document, end_url):
    html = fetch_page(url)
    if not html:
        print("Gagal memuat halaman.")
        return None

    soup = BeautifulSoup(html, 'html.parser')

    # Judul
    title_tag = soup.find('span', class_='chr-text')
    chapter_title = title_tag.get_text(strip=True) if title_tag else "No Title"

    # Konten
    content = soup.find('div', class_='chr-c')
    if not content:
        print("Konten tidak ditemukan.")
        return None

    # Hapus elemen sampah
    for bad in content.find_all(['script', 'style', 'table', 'noscript']):
        bad.decompose()

    text = content.get_text(separator="\n")
    lines = clean_text(text)

    # ---- OUTPUT ----
    print(f"Processing: {chapter_title}")

    document.add_heading(chapter_title, level=1)

    # Tambahkan paragraf rapi (tanpa double newlines)
    for line in lines:
        document.add_paragraph(line)

    # Cek stop
    if url.rstrip("/") == end_url.rstrip("/"):
        print("Reached end chapter.")
        return None

    # URL next
    next_btn = soup.find('a', id='next_chap')
    if not next_btn:
        print("Next chapter not found. Stopping.")
        return None

    next_url = urljoin(url, next_btn['href'])
    return next_url



# --- USER CONFIG ---
start_url = "https://novelbin.com/b/complete-martial-arts-attributes/chapter-1304-i-was-playing-with-you"
end_url   = "https://novelbin.com/b/complete-martial-arts-attributes/chapter-1600-choice-2"


# --- RUN ---
doc = Document()
doc.add_heading("Scraped Novel Content", level=0)

current_url = start_url

while current_url:
    print(f"Scraping {current_url}")
    current_url = scrape_chapter(current_url, doc, end_url)
    time.sleep(DELAY)

doc.save("CMAA_1304-1600.docx")
print("Scraping selesai Fedy. Happy reading!")
