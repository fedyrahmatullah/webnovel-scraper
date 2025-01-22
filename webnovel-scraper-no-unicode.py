import requests
from bs4 import BeautifulSoup
import time
from docx import Document
import re
from urllib.parse import urljoin  # Untuk menggabungkan URL relatif

# Fungsi untuk scrape konten dari setiap halaman
def scrape_chapter(url, document, end_url):
    # Kirim request ke halaman
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')

    # Ambil judul chapter
    chapter_title = soup.find('span', {'class': 'chr-text'}).get_text() if soup.find('span', {'class': 'chr-text'}) else "No Title"
    
    # Menghapus karakter newline berlebih setelah judul
    chapter_title = chapter_title.strip()  # strip() untuk menghapus spasi kosong di awal dan akhir

    # Temukan konten teks di halaman (ubah selector sesuai struktur HTML)
    content = soup.find('div', {'class': 'chr-c'})  # Ubah class sesuai HTML halaman

    if content:
        # Keluarkan teks konten
        chapter_text = content.get_text(separator='\n')

        # Hapus karakter yang tidak diinginkan (seperti ♂, &nbsp;, dan 　) bisa di cek di https://www.soscisurvey.de/tools/view-chars.php
        chapter_text = re.sub(r'♂+|&nbsp;||0b|U+3000　', '', chapter_text)

        print(f"Chapter Title: {chapter_title}")
        print(chapter_text)

        # Tambahkan judul chapter ke dokumen Word
        document.add_heading(chapter_title, level=1)

        # Tambahkan setiap baris konten ke dokumen Word dengan satu enter setelahnya
        for line in chapter_text.split('\n'):
            document.add_paragraph(line)  # Menambahkan baris
            document.add_paragraph("")  # Baris kosong setelah setiap baris teks

        # Jika sudah mencapai halaman terakhir, hentikan proses
        if url == end_url:
            print("Mencapai halaman terakhir.")
            return None

        # Temukan link ke chapter berikutnya (menggunakan ID 'next_chap')
        next_button = soup.find('a', {'id': 'next_chap'})
        if next_button:
            next_url = next_button['href']
            # Gabungkan URL relatif dengan domain dasar
            next_url = urljoin(url, next_url)
            return next_url  # Kembalikan URL halaman berikutnya
        else:
            return None  # Tidak ada halaman berikutnya
    else:
        print("Konten tidak ditemukan di halaman ini.")
        return None

# URL halaman pertama
start_url = 'https://novelbin.com/b/complete-martial-arts-attributes/chapter-1863-find-your-own-spot-and-stay-there-1'

# URL halaman terakhir (batas scraping)
end_url = 'https://novelbin.com/b/complete-martial-arts-attributes/chapter-2500-no-one-can-save-him-teach-him-a-lesson-heartache-4'

# Buat dokumen Word baru
doc = Document()
doc.add_heading('Scraped Novel Content', 0)  # Judul dokumen

# Loop untuk scrape semua chapter yang diinginkan
current_url = start_url

while current_url:
    print(f"Scraping {current_url} ...")
    current_url = scrape_chapter(current_url, doc, end_url)
    time.sleep(2)  # Tambahkan delay untuk menghindari memblokir server

# Simpan dokumen ke file .docx
doc.save('CMAA1866-2500.docx')

print("Scraping selesai fedy. selamat membaca")
