from docx import Document

# Fungsi untuk mendeteksi dan menghapus paragraf kosong berturut-turut
def remove_multiple_empty_paragraphs(doc, threshold=10):
    empty_count = 0  # Menghitung jumlah paragraf kosong berturut-turut
    paragraphs_to_delete = []  # Menyimpan indeks paragraf yang akan dihapus

    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == "":  # Paragraf kosong
            empty_count += 1
        else:
            if empty_count >= threshold:  # Jika ditemukan 10+ paragraf kosong berturut-turut
                paragraphs_to_delete.extend(range(i - empty_count, i))
            empty_count = 0  # Reset hitungan untuk paragraf kosong berikutnya

    # Hapus paragraf yang ditandai dari akhir ke awal untuk menghindari indeks bergeser
    for idx in reversed(paragraphs_to_delete):
        p = doc.paragraphs[idx]._element
        p.getparent().remove(p)

# Muat dokumen Word
input_file = 'CMAA1304-1600cleaneda.docx'
doc = Document(input_file)

# Hapus paragraf kosong berturut-turut
remove_multiple_empty_paragraphs(doc, threshold=10)

# Simpan hasilnya ke file baru
output_file = 'CMAA1304-1600cleaned_no_empty_paragraphs55.docx'
doc.save(output_file)

print(f"Proses selesai. Dokumen disimpan sebagai '{output_file}'.")
